import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

TEMPLATE_PATH = Path(__file__).parent.parent / "template" / "template.docx"
ASSETS_PATH = Path(__file__).parent.parent / "assets"


def _safe_add_paragraph(doc, text, *style_names):
    """Add a paragraph unstyled first, then apply the first working style."""
    para = doc.add_paragraph(text)
    for name in style_names:
        try:
            para.style = doc.styles[name]
            return para
        except (KeyError, Exception):
            continue
    return para


def _safe_style(doc, *names):
    """Return the first style name that works in the document, or None for default."""
    for name in names:
        try:
            doc.styles.get_style_id(name, __import__("docx.enum.style", fromlist=["WD_STYLE_TYPE"]).WD_STYLE_TYPE.PARAGRAPH)
            return name
        except (KeyError, Exception):
            continue
    return None


def _replace_para_text(para, new_text):
    for run in para.runs:
        run._element.getparent().remove(run._element)
    para.add_run(new_text)


def _set_cell_shading(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tcPr.append(shd)


def _set_cell_border(cell, border_hex="CDD3D6"):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), border_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_screenshot_placeholder(doc, n, page_name):
    table = doc.add_table(rows=1, cols=1)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "EBEBEB")
    _set_cell_border(cell, "CDD3D6")

    trPr = table.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(8 * 567)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run(f"<<<Paste Screenshot {n}: {page_name}>>>")
    run1.font.size = Pt(14)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        f'(Insert a screenshot of the "{page_name}" page from the published Power BI report)'
    )
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _set_page_break_before(para):
    pPr = para._element.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pPr.insert(0, pb)


def _parse_md_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        if re.match(r"^\|[-| :]+\|$", stripped):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_word_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data[:max_cols]):
            cell = row.cells[j]
            para = cell.paragraphs[0]
            style_name = "Table Heading" if i == 0 else "Table Text"
            try:
                para.style = doc.styles[style_name]
            except KeyError:
                pass
            para.text = cell_text


def _strip_inline_md(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _inject_content(doc, content_md, assets_path=None):
    lines = content_md.split("\n")
    screenshot_counter = 0
    i = 0

    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            para = _safe_add_paragraph(doc, _strip_inline_md(line[2:].strip()), "Heading 1")
            if not para.style or "Heading" not in para.style.name:
                for run in para.runs:
                    run.bold = True; run.font.size = Pt(16)
            _set_page_break_before(para)

        elif line.startswith("## "):
            _safe_add_paragraph(doc, _strip_inline_md(line[3:].strip()), "Heading 2")

        elif line.startswith("### "):
            _safe_add_paragraph(doc, _strip_inline_md(line[4:].strip()), "Heading 3", "Heading 2")

        elif line.startswith("#### "):
            _safe_add_paragraph(doc, _strip_inline_md(line[5:].strip()), "Heading 4", "Heading 3")

        elif line.startswith("- "):
            _safe_add_paragraph(doc, _strip_inline_md(line[2:].strip()), "List Bullet", "List Paragraph")

        elif re.match(r"^\d+\. ", line):
            _safe_add_paragraph(doc, _strip_inline_md(re.sub(r"^\d+\. ", "", line)), "List Number", "List Paragraph")

        elif re.match(r"^<<<BOILERPLATE:\s*.+>>>$", line.strip()):
            filename = re.sub(r"^<<<BOILERPLATE:\s*", "", line.strip()).rstrip(">").strip()
            resolved_assets = assets_path if assets_path is not None else ASSETS_PATH
            boilerplate_path = resolved_assets / filename
            if not boilerplate_path.exists():
                raise FileNotFoundError(
                    f"Boilerplate file not found: {boilerplate_path}"
                )
            _inject_content(doc, boilerplate_path.read_text(encoding="utf-8-sig"), assets_path=resolved_assets)

        elif line.strip().startswith("<<<SCREENSHOT:"):
            name = re.sub(r"^<<<SCREENSHOT:\s*", "", line.strip()).rstrip(">").strip()
            screenshot_counter += 1
            _add_screenshot_placeholder(doc, screenshot_counter, name)

        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_word_table(doc, _parse_md_table(table_lines))
            continue

        elif line.strip():
            text = _strip_inline_md(line.strip())
            next_line = lines[i + 1] if i + 1 < len(lines) else ""
            if line.strip().endswith(":") and (
                next_line.startswith("- ") or re.match(r"^\d+\. ", next_line)
            ):
                p = _safe_add_paragraph(doc, "", "Normal")
                p.add_run(text).bold = True
            else:
                _safe_add_paragraph(doc, text, "Normal")

        i += 1


def _update_fields_via_word(docx_path):
    import win32com.client
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    resolved = str(Path(docx_path).resolve())
    try:
        doc = word.Documents.Open(resolved)
        doc.Fields.Update()
        if doc.TablesOfContents.Count > 0:
            doc.TablesOfContents(1).Update()
        if doc.ReadOnly:
            doc.SaveAs2(resolved)
        else:
            doc.Save()
        doc.Close()
    finally:
        word.Quit()


def build_docx(content_md, metadata, output_path, template_path=None, assets_path=None):
    resolved_template = Path(template_path) if template_path is not None else TEMPLATE_PATH
    doc = Document(str(resolved_template))
    raw_name = metadata.get("name", "Dashboard")
    stripped = re.sub(r"[_ ]?Dashboard$", "", raw_name, flags=re.IGNORECASE).strip()
    if stripped:
        display_name = f"{stripped} Dashboard"
    else:
        display_name = raw_name

    d = datetime.now()
    date_str = f"{d.day} {d.strftime('%B %Y')}"

    doc.core_properties.title = f"User Guide - {display_name} (Power BI User Guide)"

    # Replace named placeholders (works with any template structure)
    for p in doc.paragraphs:
        if "{Dashboard Name}" in p.text:
            _replace_para_text(p, p.text.replace("{Dashboard Name}", display_name))
        elif "{date}" in p.text:
            _replace_para_text(p, p.text.replace("{date}", date_str))
        elif re.match(r"^Version:\s*\d", p.text):
            _replace_para_text(p, "Version: 1.0")

    body = doc.element.body
    paras = list(doc.paragraphs)
    for p in paras[11:]:
        body.remove(p._element)

    # Remove orphaned non-paragraph body elements (tables, SDTs) after the trim
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    kept = {p._element for p in doc.paragraphs}
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag not in ("p", "sectPr") and child not in kept:
            body.remove(child)

    _inject_content(doc, content_md, assets_path=assets_path)

    end_para = _safe_add_paragraph(doc, "End of User Guide", "Normal")
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in end_para.runs:
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(output_path))
    print(f"Saved: {output_path}", file=sys.stderr)
    _update_fields_via_word(output_path)
    print(f"Fields updated via Word: {output_path}", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description="Write branded DOCX user guide")
    parser.add_argument("--content", required=True, help="Path to markdown content file")
    parser.add_argument("--metadata", required=True, help="Path to metadata JSON file")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    parser.add_argument("--template", default=None, help="Path to .docx template (default: assets/template.docx)")
    parser.add_argument("--assets", default=None, help="Path to assets directory containing boilerplate .md files (default: assets/)")
    args = parser.parse_args()

    content_path = Path(args.content)
    metadata_path = Path(args.metadata)
    output_path = Path(args.output)
    template_path = Path(args.template) if args.template else None
    assets_path = Path(args.assets) if args.assets else None

    resolved_template = template_path if template_path is not None else TEMPLATE_PATH

    if not content_path.exists():
        print(f"Error: Content file not found: {content_path}", file=sys.stderr)
        sys.exit(1)
    if not metadata_path.exists():
        print(f"Error: Metadata file not found: {metadata_path}", file=sys.stderr)
        sys.exit(1)
    if not resolved_template.exists():
        print(f"Error: Template not found: {resolved_template}", file=sys.stderr)
        print("Fix: Ensure template.docx exists or pass --template <path>", file=sys.stderr)
        sys.exit(1)

    content_md = content_path.read_text(encoding="utf-8-sig")
    metadata = json.loads(metadata_path.read_text(encoding="utf-8-sig"))

    print(f"Building DOCX for {metadata.get('name', 'unknown')}...", file=sys.stderr)
    build_docx(content_md, metadata, output_path, template_path=template_path, assets_path=assets_path)


if __name__ == "__main__":
    main()
