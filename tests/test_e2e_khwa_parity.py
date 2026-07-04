import json
import subprocess
import sys
from pathlib import Path

import pytest

SCRIPTS = Path(__file__).parent.parent / "scripts"
PBIP = Path("C:/Users/Admin/Documents/Life/projects/pbi-dashboard-factory/output/Finance_Dashboard/Finance_Dashboard.pbip")

sys.path.insert(0, str(SCRIPTS))
from write_docx import _inject_content, TEMPLATE_PATH
from docx import Document


def _parse_finance():
    result = subprocess.run(
        [sys.executable, str(SCRIPTS / "parse_pbip.py"), str(PBIP)],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, f"parse_pbip.py failed: {result.stderr}"
    return json.loads(result.stdout)


def _make_doc():
    return Document(str(TEMPLATE_PATH))


def test_finance_parse_pbip_has_sources():
    data = _parse_finance()
    assert "sources" in data, "Expected 'sources' key in parse output"
    assert isinstance(data["sources"], list), "Expected 'sources' to be a list"
    assert len(data["sources"]) > 0, "Expected 'sources' to be a non-empty list"


def test_finance_parse_pbip_has_acronyms():
    data = _parse_finance()
    assert "acronyms" in data, "Expected 'acronyms' key in parse output"
    assert isinstance(data["acronyms"], list), "Expected 'acronyms' to be a list"


def test_finance_parse_pbip_has_kpi_visuals():
    data = _parse_finance()
    assert "kpi_visuals" in data, "Expected 'kpi_visuals' key in parse output"
    assert isinstance(data["kpi_visuals"], dict), "Expected 'kpi_visuals' to be a dict"


def test_finance_kpi_visuals_keyed_by_pages():
    data = _parse_finance()
    page_names = {p["displayName"] for p in data["report"]["pages"]}
    kpi_keys = set(data["kpi_visuals"].keys())
    assert kpi_keys == page_names, (
        f"kpi_visuals keys {kpi_keys} do not match page displayNames {page_names}"
    )


def test_finance_sources_have_required_fields():
    data = _parse_finance()
    for i, source in enumerate(data["sources"]):
        for field in ("table", "type", "query_snippet"):
            assert field in source, (
                f"Source[{i}] missing required field '{field}': {source}"
            )


def test_finance_acronyms_have_required_fields():
    data = _parse_finance()
    for i, acr in enumerate(data["acronyms"]):
        for field in ("acronym", "found_in"):
            assert field in acr, (
                f"Acronym[{i}] missing required field '{field}': {acr}"
            )


def _doc_full_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(parts)


def test_all_boilerplate_files_inject_correctly():
    cases = [
        ("boilerplate_dashboard_functions.md", "Navigation Pane"),
        ("boilerplate_tips_and_tools.md", None),
        ("boilerplate_export.md", "Export"),
        ("boilerplate_access_form.md", "SARA"),
    ]
    for filename, expected_text in cases:
        doc = _make_doc()
        _inject_content(doc, f"<<<BOILERPLATE: {filename}>>>")
        combined = _doc_full_text(doc)
        if filename == "boilerplate_tips_and_tools.md":
            assert "Tooltip" in combined or "tooltip" in combined, (
                f"{filename}: expected 'Tooltip'/'tooltip' in content, got: {combined[:300]}"
            )
        else:
            assert expected_text in combined, (
                f"{filename}: expected '{expected_text}' in content, got: {combined[:300]}"
            )


def test_expected_section_headings_present():
    doc = _make_doc()
    content = "\n".join([
        "# Background",
        "",
        "# Document Purpose",
        "",
        "<<<BOILERPLATE: boilerplate_dashboard_functions.md>>>",
        "",
        "# How to Access the Dashboard",
        "",
        "<<<BOILERPLATE: boilerplate_tips_and_tools.md>>>",
        "",
        "<<<BOILERPLATE: boilerplate_export.md>>>",
        "",
        "# APPENDIX",
    ])
    _inject_content(doc, content)
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    expected_headings = [
        "Background",
        "Document Purpose",
        "Dashboard Functions",
        "How to Access the Dashboard",
        "Other Useful Tips and Tools",
        "Export (Excel / PDF)",
        "APPENDIX",
    ]
    for heading in expected_headings:
        assert heading in heading_texts, (
            f"Expected heading '{heading}' not found. Headings present: {heading_texts}"
        )
