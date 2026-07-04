import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))
from write_docx import _inject_content, build_docx, TEMPLATE_PATH
from docx import Document

OUTPUT_DIR = Path(__file__).parent.parent / "output"
TEMPLATE_DIR = Path(__file__).parent.parent / "template"


def _make_doc():
    return Document(str(TEMPLATE_PATH))


def test_boilerplate_injection_inserts_content():
    doc = _make_doc()
    content = "# Test\n<<<BOILERPLATE: boilerplate_dashboard_functions.md>>>\n"
    _inject_content(doc, content)
    texts = [p.text for p in doc.paragraphs]
    combined = " ".join(texts)
    assert "Dashboard Functions" in combined or "Navigation Pane" in combined, (
        f"Expected boilerplate content in paragraphs, got: {texts}"
    )


def test_boilerplate_injection_unknown_file_raises():
    doc = _make_doc()
    content = "<<<BOILERPLATE: nonexistent_file.md>>>"
    with pytest.raises(FileNotFoundError):
        _inject_content(doc, content)


def test_heading4_before_bullets_uses_bold_normal():
    doc = _make_doc()
    content = "Filters available:\n- Filter A\n- Filter B\n"
    _inject_content(doc, content)
    intro_para = next(
        (p for p in doc.paragraphs if p.text == "Filters available:"),
        None,
    )
    assert intro_para is not None, "Paragraph 'Filters available:' not found"
    # Style may be None (no explicit style = document default) when template lacks 'Normal'
    style_name = intro_para.style.name if intro_para.style else ""
    assert "Normal" in style_name or intro_para.style is None, (
        f"Expected Normal or unstyled, got: {style_name}"
    )
    assert any(run.bold for run in intro_para.runs), "Expected bold run in intro paragraph"
    bullet_texts = [p.text for p in doc.paragraphs if p.style and "List" in p.style.name]
    assert "Filter A" in bullet_texts, f"Expected 'Filter A' in list bullets, got: {bullet_texts}"


def test_title_no_double_dashboard(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    metadata = {"name": "Finance_Dashboard"}
    out = tmp_path / "test.docx"
    build_docx("# Test", metadata, out)
    result_doc = Document(str(out))
    all_texts = [p.text for p in result_doc.paragraphs]
    assert not any("Dashboard Dashboard" in t for t in all_texts), (
        f"'Dashboard Dashboard' found in: {all_texts}"
    )
    assert any("Finance Dashboard" in t for t in all_texts), (
        f"Display name 'Finance Dashboard' not found in: {all_texts}"
    )


def test_title_edge_case_dashboard_only_name(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    metadata = {"name": "Dashboard"}
    out = tmp_path / "test.docx"
    build_docx("# Test", metadata, out)
    result_doc = Document(str(out))
    all_texts = [p.text for p in result_doc.paragraphs]
    assert not any("Dashboard Dashboard" in t for t in all_texts), (
        f"'Dashboard Dashboard' found in: {all_texts}"
    )


def test_placeholder_dashboard_name_replaced(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    out = tmp_path / "out.docx"
    build_docx("# Hello", {"name": "Test Dashboard"}, out)
    doc = Document(str(out))
    all_texts = [p.text for p in doc.paragraphs]
    assert not any("{Dashboard Name}" in t for t in all_texts), (
        f"Unreplaced placeholder '{{Dashboard Name}}' found: {all_texts}"
    )
    assert any("Test Dashboard" in t for t in all_texts), (
        f"Display name not found in: {all_texts}"
    )


def test_placeholder_date_replaced(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    out = tmp_path / "out.docx"
    build_docx("# Hello", {"name": "My Report"}, out)
    doc = Document(str(out))
    all_texts = [p.text for p in doc.paragraphs]
    assert not any("{date}" in t for t in all_texts), (
        f"Unreplaced placeholder '{{date}}' found: {all_texts}"
    )


def test_no_orphan_tables_after_trim(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    out = tmp_path / "out.docx"
    build_docx("# Hello", {"name": "Test"}, out)
    doc = Document(str(out))
    assert len(doc.tables) == 0, (
        f"Orphan tables remain in output: {len(doc.tables)} table(s)"
    )


def test_build_docx_custom_template(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    custom_template = TEMPLATE_PATH
    out = tmp_path / "out.docx"
    build_docx("# Hello", {"name": "Test"}, out, template_path=custom_template)
    assert out.exists()


def test_build_docx_custom_assets(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    assets = Path(__file__).parent.parent / "assets"
    out = tmp_path / "out.docx"
    build_docx(
        "# Test\n<<<BOILERPLATE: boilerplate_dashboard_functions.md>>>",
        {"name": "Test"},
        out,
        assets_path=assets,
    )
    result_doc = Document(str(out))
    texts = " ".join(p.text for p in result_doc.paragraphs)
    assert "Dashboard Functions" in texts or "Navigation Pane" in texts


def test_build_docx_missing_custom_template_raises(tmp_path, monkeypatch):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    out = tmp_path / "out.docx"
    with pytest.raises(Exception):
        build_docx("# Test", {"name": "Test"}, out, template_path=tmp_path / "no_such.docx")


def test_inject_content_custom_assets(tmp_path):
    assets = tmp_path / "assets"
    assets.mkdir()
    (assets / "custom_boilerplate.md").write_text("Custom boilerplate text")
    doc = Document(str(TEMPLATE_PATH))
    _inject_content(doc, "<<<BOILERPLATE: custom_boilerplate.md>>>", assets_path=assets)
    texts = [p.text for p in doc.paragraphs]
    assert any("Custom boilerplate text" in t for t in texts)


def test_cli_template_flag(tmp_path, monkeypatch):
    import subprocess, sys
    script = Path(__file__).parent.parent / "scripts" / "write_docx.py"
    content_file = tmp_path / "content.md"
    meta_file = tmp_path / "meta.json"
    out_file = tmp_path / "out.docx"
    content_file.write_text("# Hello")
    meta_file.write_text('{"name": "CLI Test"}')
    result = subprocess.run(
        [
            sys.executable, str(script),
            "--content", str(content_file),
            "--metadata", str(meta_file),
            "--output", str(out_file),
            "--template", str(TEMPLATE_PATH),
            "--assets", str(Path(__file__).parent.parent / "assets"),
        ],
        capture_output=True, text=True
    )
    assert result.returncode == 0, f"CLI failed: {result.stderr}"
    assert out_file.exists()


# --- Path structure tests (TDD: write first, implement after) ---

def test_template_path_is_in_template_dir():
    assert TEMPLATE_PATH.parent.name == "template", (
        f"TEMPLATE_PATH should be in 'template/' dir, got '{TEMPLATE_PATH.parent.name}'"
    )


def test_template_path_filename():
    assert TEMPLATE_PATH.name == "template.docx", (
        f"Expected template.docx, got {TEMPLATE_PATH.name}"
    )


def test_template_path_exists():
    assert TEMPLATE_PATH.exists(), f"Template not found at {TEMPLATE_PATH}"


def test_output_dir_exists():
    assert OUTPUT_DIR.exists(), f"output/ directory missing at {OUTPUT_DIR}"


def test_build_docx_writes_to_output_dir(monkeypatch, tmp_path):
    monkeypatch.setattr("write_docx._update_fields_via_word", lambda path: None)
    out = OUTPUT_DIR / "test_output.docx"
    build_docx("# Test Output Dir", {"name": "Test"}, out)
    assert out.exists()
    out.unlink()
